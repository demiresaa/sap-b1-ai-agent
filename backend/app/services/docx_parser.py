"""DOCX sipariş belgesi parser'ı.

python-docx ile paragraf ve tablo içeriğini düz metin + TSV formatında çıkarır.
Çıktı, `document_reader._extract_via_text()` ile aynı LLM akışına verilir.
"""
from __future__ import annotations

import logging
import pathlib
from dataclasses import dataclass, field

import docx

logger = logging.getLogger(__name__)

MAX_TABLE_ROWS = 100


@dataclass(slots=True)
class DocxContent:
    """Normalize edilmiş DOCX içeriği."""

    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)

    def to_prompt_text(self) -> str:
        """LLM prompt'una gönderilebilir düz metin + TSV tablo temsili."""
        parts: list[str] = []

        if self.paragraphs:
            parts.append("\n".join(p for p in self.paragraphs if p.strip()))

        for i, table in enumerate(self.tables, start=1):
            if not table:
                continue
            parts.append(f"\n--- Tablo {i} ---")
            for row in table[:MAX_TABLE_ROWS]:
                parts.append(" | ".join(cell for cell in row))
            if len(table) > MAX_TABLE_ROWS:
                parts.append(f"… ve {len(table) - MAX_TABLE_ROWS} satır daha")

        return "\n".join(parts)


def parse_docx(path: pathlib.Path | str) -> DocxContent:
    """DOCX dosyasını okuyup `DocxContent` döner.

    Paragrafları sırasıyla, tabloları TSV benzeri satırlar olarak toplar.
    """
    p = pathlib.Path(path)
    if not p.exists():
        raise FileNotFoundError(f"DOCX bulunamadı: {p}")

    document = docx.Document(str(p))

    paragraphs: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    tables: list[list[list[str]]] = []
    for tbl in document.tables:
        rows: list[list[str]] = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(c for c in cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    logger.info(
        "[docx] dosya=%s paragraf=%d tablo=%d",
        p.name,
        len(paragraphs),
        len(tables),
    )
    return DocxContent(paragraphs=paragraphs, tables=tables)
